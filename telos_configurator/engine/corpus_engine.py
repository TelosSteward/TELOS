"""
TELOS Corpus Engine - Document management and embedding system.

Handles document loading, parsing, embedding generation, and semantic search
for the TELOS Corpus Configurator MVP.
"""

import io
import json
import uuid
import logging
import threading
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple, Callable
from dataclasses import dataclass, asdict

import numpy as np
import requests


# Configure logging
logger = logging.getLogger(__name__)


@dataclass
class CorpusDocument:
    """
    Represents a document in the corpus.

    Attributes:
        doc_id: Unique identifier for the document
        filename: Original filename
        title: Document title (extracted or from metadata)
        category: Document category/type
        source: Source system or origin
        text_content: Full text content of the document
        key_provisions: List of important provisions or sections
        escalation_triggers: List of conditions requiring escalation
        embedding: NumPy array of embedding vector (None if not embedded)
        embedded: Flag indicating if document has been embedded
        added_at: ISO timestamp of when document was added
    """
    doc_id: str
    filename: str
    title: str
    category: str
    source: str
    text_content: str
    key_provisions: List[str]
    escalation_triggers: List[str]
    embedding: Optional[np.ndarray]
    embedded: bool
    added_at: str


class CorpusEngine:
    """
    Core engine for corpus document management and semantic search.

    Features:
    - Document loading from multiple formats (JSON, PDF, TXT, MD, DOCX, XLSX)
    - Embedding generation using Ollama nomic-embed-text
    - Corpus persistence (save/load)
    - Semantic search using cosine similarity
    - Thread-safe operations for Streamlit
    """

    # Maximum file size: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024

    # Ollama embedding configuration
    OLLAMA_ENDPOINT = "http://localhost:11434/api/embeddings"
    EMBEDDING_MODEL = "nomic-embed-text"
    EMBEDDING_TIMEOUT = 30

    def __init__(self):
        """Initialize the corpus engine."""
        self.documents: Dict[str, CorpusDocument] = {}
        self._lock = threading.Lock()
        logger.info("CorpusEngine initialized")

    # ========== Document Loading ==========

    def add_document(
        self,
        uploaded_file,
        category: str = "general",
        source: str = "upload"
    ) -> Tuple[bool, str, Optional[str]]:
        """
        Add a document to the corpus from a Streamlit UploadedFile.

        Args:
            uploaded_file: Streamlit UploadedFile object
            category: Document category (default: "general")
            source: Document source (default: "upload")

        Returns:
            Tuple of (success, message, doc_id)
            - success: True if document added successfully
            - message: Status/error message
            - doc_id: Document ID if successful, None otherwise
        """
        try:
            # Check file size
            file_size = uploaded_file.size
            if file_size > self.MAX_FILE_SIZE:
                return (
                    False,
                    f"File too large ({file_size / 1024 / 1024:.1f}MB). Max size is 10MB.",
                    None
                )

            # Read file bytes
            file_bytes = uploaded_file.read()
            filename = uploaded_file.name
            file_type = uploaded_file.type

            # Extract content based on file type
            content, error = self._extract_content(file_bytes, filename, file_type)

            if error:
                return False, error, None

            if not content:
                return False, "No text content could be extracted from file", None

            # Parse document structure
            doc_data = self._parse_document_structure(content, filename)

            # Create document
            doc_id = str(uuid.uuid4())
            document = CorpusDocument(
                doc_id=doc_id,
                filename=filename,
                title=doc_data.get("title", filename),
                category=category,
                source=source,
                text_content=doc_data["text_content"],
                key_provisions=doc_data.get("key_provisions", []),
                escalation_triggers=doc_data.get("escalation_triggers", []),
                embedding=None,
                embedded=False,
                added_at=datetime.utcnow().isoformat()
            )

            # Add to corpus (thread-safe)
            with self._lock:
                self.documents[doc_id] = document

            logger.info(f"Added document: {filename} (ID: {doc_id})")
            return True, f"Successfully added: {filename}", doc_id

        except Exception as e:
            error_msg = f"Error adding document: {str(e)}"
            logger.error(error_msg)
            return False, error_msg, None

    def _extract_content(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: str
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text content from file bytes.

        Args:
            file_bytes: Raw file content
            filename: Original filename
            file_type: MIME type

        Returns:
            Tuple of (content, error_message)
        """
        try:
            ext = Path(filename).suffix.lower().lstrip('.')

            # JSON files
            if ext == 'json' or file_type == 'application/json':
                return self._extract_json(file_bytes)

            # PDF files
            if ext == 'pdf' or file_type == 'application/pdf':
                return self._extract_pdf(file_bytes)

            # Text-based files (TXT, MD)
            if ext in ['txt', 'md'] or file_type in ['text/plain', 'text/markdown']:
                return self._extract_text(file_bytes)

            # DOCX files
            if ext == 'docx':
                return self._extract_docx(file_bytes)

            # XLSX files
            if ext == 'xlsx':
                return self._extract_xlsx(file_bytes)

            return None, f"Unsupported file type: {file_type}"

        except Exception as e:
            return None, f"Error extracting content: {str(e)}"

    def _extract_json(self, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract content from JSON file."""
        try:
            content = file_bytes.decode('utf-8')
            # Validate JSON
            json.loads(content)
            return content, None
        except json.JSONDecodeError as e:
            return None, f"Invalid JSON: {str(e)}"
        except UnicodeDecodeError:
            return None, "Could not decode JSON file"

    def _extract_text(self, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract content from text files."""
        try:
            return file_bytes.decode('utf-8'), None
        except UnicodeDecodeError:
            try:
                return file_bytes.decode('latin-1'), None
            except Exception as e:
                return None, f"Could not decode text file: {str(e)}"

    def _extract_pdf(self, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            if not text_parts:
                return None, "No text could be extracted from PDF"

            return "\n\n".join(text_parts), None

        except ImportError:
            return None, "PDF support requires PyPDF2. Install with: pip install PyPDF2"
        except Exception as e:
            return None, f"Error extracting PDF: {str(e)}"

    def _extract_docx(self, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from DOCX file."""
        try:
            import docx
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            if not text_parts:
                return None, "No text could be extracted from DOCX"

            return "\n\n".join(text_parts), None

        except ImportError:
            return None, "DOCX support requires python-docx. Install with: pip install python-docx"
        except Exception as e:
            return None, f"Error extracting DOCX: {str(e)}"

    def _extract_xlsx(self, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
        """Extract text from XLSX file."""
        try:
            import openpyxl
            xlsx_file = io.BytesIO(file_bytes)
            workbook = openpyxl.load_workbook(xlsx_file, read_only=True)

            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}\n")

                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)

                text_parts.append("\n")

            workbook.close()

            if not text_parts:
                return None, "No text could be extracted from XLSX"

            return "\n".join(text_parts), None

        except ImportError:
            return None, "XLSX support requires openpyxl. Install with: pip install openpyxl"
        except Exception as e:
            return None, f"Error extracting XLSX: {str(e)}"

    def _parse_document_structure(self, content: str, filename: str) -> Dict[str, Any]:
        """
        Parse document content to extract structured information.

        For JSON files, attempts to extract fields like:
        - title, text_content, key_provisions, escalation_triggers

        For other formats, treats entire content as text_content.

        Args:
            content: Raw text content
            filename: Original filename

        Returns:
            Dictionary with parsed fields
        """
        # Try to parse as JSON
        try:
            data = json.loads(content)

            # Extract fields (flexible schema)
            result = {
                "title": data.get("title", data.get("name", filename)),
                "text_content": data.get("text_content", data.get("content", content)),
                "key_provisions": data.get("key_provisions", []),
                "escalation_triggers": data.get("escalation_triggers", [])
            }

            # Ensure lists
            if not isinstance(result["key_provisions"], list):
                result["key_provisions"] = []
            if not isinstance(result["escalation_triggers"], list):
                result["escalation_triggers"] = []

            return result

        except json.JSONDecodeError:
            # Not JSON, treat as plain text
            return {
                "title": filename,
                "text_content": content,
                "key_provisions": [],
                "escalation_triggers": []
            }

    # ========== Document Management ==========

    def remove_document(self, doc_id: str) -> Tuple[bool, str]:
        """
        Remove a document from the corpus.

        Args:
            doc_id: Document ID to remove

        Returns:
            Tuple of (success, message)
        """
        with self._lock:
            if doc_id not in self.documents:
                return False, f"Document not found: {doc_id}"

            doc = self.documents.pop(doc_id)
            logger.info(f"Removed document: {doc.filename} (ID: {doc_id})")
            return True, f"Removed: {doc.filename}"

    def list_documents(self) -> List[Dict[str, Any]]:
        """
        Get list of all documents with metadata.

        Returns:
            List of document metadata dictionaries
        """
        with self._lock:
            result = []
            for doc in self.documents.values():
                result.append({
                    "doc_id": doc.doc_id,
                    "filename": doc.filename,
                    "title": doc.title,
                    "category": doc.category,
                    "source": doc.source,
                    "embedded": doc.embedded,
                    "added_at": doc.added_at,
                    "text_length": len(doc.text_content),
                    "key_provisions_count": len(doc.key_provisions),
                    "escalation_triggers_count": len(doc.escalation_triggers)
                })
            return result

    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """
        Get full document data.

        Args:
            doc_id: Document ID

        Returns:
            Document dictionary or None if not found
        """
        with self._lock:
            if doc_id not in self.documents:
                return None

            doc = self.documents[doc_id]
            return {
                "doc_id": doc.doc_id,
                "filename": doc.filename,
                "title": doc.title,
                "category": doc.category,
                "source": doc.source,
                "text_content": doc.text_content,
                "key_provisions": doc.key_provisions,
                "escalation_triggers": doc.escalation_triggers,
                "embedded": doc.embedded,
                "added_at": doc.added_at
            }

    def clear_corpus(self) -> Tuple[bool, str]:
        """
        Remove all documents from the corpus.

        Returns:
            Tuple of (success, message)
        """
        with self._lock:
            count = len(self.documents)
            self.documents.clear()
            logger.info(f"Cleared corpus ({count} documents removed)")
            return True, f"Cleared {count} documents from corpus"

    # ========== Embedding Generation ==========

    def embed_document(
        self,
        doc_id: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> Tuple[bool, str]:
        """
        Generate embedding for a single document.

        Args:
            doc_id: Document ID
            progress_callback: Optional callback function for progress updates

        Returns:
            Tuple of (success, message)
        """
        with self._lock:
            if doc_id not in self.documents:
                return False, f"Document not found: {doc_id}"

            doc = self.documents[doc_id]

        if progress_callback:
            progress_callback(f"Embedding: {doc.filename}")

        # Generate embedding
        embedding = self._get_embedding(doc.text_content)

        if embedding is None:
            return False, f"Failed to generate embedding for: {doc.filename}"

        # Update document
        with self._lock:
            doc.embedding = embedding
            doc.embedded = True

        logger.info(f"Embedded document: {doc.filename} (ID: {doc_id})")
        return True, f"Embedded: {doc.filename}"

    def embed_all(
        self,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> Tuple[int, int, List[str]]:
        """
        Embed all documents in the corpus.

        Args:
            progress_callback: Optional callback(current, total, filename)

        Returns:
            Tuple of (success_count, failure_count, failed_filenames)
        """
        with self._lock:
            doc_list = list(self.documents.values())

        success_count = 0
        failure_count = 0
        failed_filenames = []
        total = len(doc_list)

        for idx, doc in enumerate(doc_list, 1):
            if progress_callback:
                progress_callback(idx, total, doc.filename)

            success, message = self.embed_document(doc.doc_id)

            if success:
                success_count += 1
            else:
                failure_count += 1
                failed_filenames.append(doc.filename)

        logger.info(f"Batch embedding complete: {success_count} success, {failure_count} failed")
        return success_count, failure_count, failed_filenames

    def get_embedding_status(self) -> Dict[str, Any]:
        """
        Get embedding status for the corpus.

        Returns:
            Dictionary with embedding statistics
        """
        with self._lock:
            total = len(self.documents)
            embedded = sum(1 for doc in self.documents.values() if doc.embedded)
            not_embedded = total - embedded

            return {
                "total_documents": total,
                "embedded": embedded,
                "not_embedded": not_embedded,
                "percentage": (embedded / total * 100) if total > 0 else 0
            }

    def _get_embedding(self, text: str) -> Optional[np.ndarray]:
        """
        Generate embedding vector for text using Ollama.

        Args:
            text: Text to embed

        Returns:
            NumPy array of embedding vector, or None on failure
        """
        try:
            response = requests.post(
                self.OLLAMA_ENDPOINT,
                json={
                    "model": self.EMBEDDING_MODEL,
                    "prompt": text
                },
                timeout=self.EMBEDDING_TIMEOUT
            )

            if response.status_code == 200:
                embedding_list = response.json().get("embedding", [])
                if embedding_list:
                    return np.array(embedding_list, dtype=np.float32)
            else:
                logger.error(f"Ollama API error: {response.status_code} - {response.text}")

            return None

        except requests.exceptions.Timeout:
            logger.error("Ollama API timeout")
            return None
        except requests.exceptions.ConnectionError:
            logger.error("Cannot connect to Ollama API. Is Ollama running?")
            return None
        except Exception as e:
            logger.error(f"Error generating embedding: {str(e)}")
            return None

    # ========== Search/Retrieval ==========

    def search(
        self,
        query_text: str,
        top_k: int = 3
    ) -> List[Dict[str, Any]]:
        """
        Search corpus for documents similar to query.

        Args:
            query_text: Search query
            top_k: Number of results to return

        Returns:
            List of result dictionaries with doc_id, similarity, and metadata
        """
        # Generate query embedding
        query_embedding = self._get_embedding(query_text)

        if query_embedding is None:
            logger.error("Failed to generate query embedding")
            return []

        return self.search_by_embedding(query_embedding, top_k)

    def search_by_embedding(
        self,
        query_embedding: np.ndarray,
        top_k: int = 3
    ) -> List[Dict[str, Any]]:
        """
        Search corpus using a pre-computed query embedding.

        Args:
            query_embedding: Query embedding vector
            top_k: Number of results to return

        Returns:
            List of result dictionaries sorted by similarity (descending)
        """
        with self._lock:
            embedded_docs = [
                doc for doc in self.documents.values()
                if doc.embedded and doc.embedding is not None
            ]

        if not embedded_docs:
            logger.warning("No embedded documents available for search")
            return []

        # Calculate similarities
        results = []
        for doc in embedded_docs:
            similarity = self._cosine_similarity(query_embedding, doc.embedding)
            results.append({
                "doc_id": doc.doc_id,
                "filename": doc.filename,
                "title": doc.title,
                "category": doc.category,
                "source": doc.source,
                "similarity": float(similarity),
                "text_preview": doc.text_content[:200] + "..." if len(doc.text_content) > 200 else doc.text_content,
                "key_provisions": doc.key_provisions,
                "escalation_triggers": doc.escalation_triggers
            })

        # Sort by similarity (descending) and limit to top_k
        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results[:top_k]

    def _cosine_similarity(self, vec1: np.ndarray, vec2: np.ndarray) -> float:
        """
        Calculate cosine similarity between two vectors.

        Args:
            vec1: First vector
            vec2: Second vector

        Returns:
            Cosine similarity score (0-1)
        """
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)

        if norm1 == 0 or norm2 == 0:
            return 0.0

        return float(dot_product / (norm1 * norm2))

    # ========== Corpus Persistence ==========

    def save_corpus(self, filepath: str) -> Tuple[bool, str]:
        """
        Save corpus to JSON file.

        Args:
            filepath: Path to save file

        Returns:
            Tuple of (success, message)
        """
        try:
            with self._lock:
                corpus_data = {
                    "version": "1.0",
                    "saved_at": datetime.utcnow().isoformat(),
                    "document_count": len(self.documents),
                    "documents": []
                }

                for doc in self.documents.values():
                    doc_dict = {
                        "doc_id": doc.doc_id,
                        "filename": doc.filename,
                        "title": doc.title,
                        "category": doc.category,
                        "source": doc.source,
                        "text_content": doc.text_content,
                        "key_provisions": doc.key_provisions,
                        "escalation_triggers": doc.escalation_triggers,
                        "embedding": doc.embedding.tolist() if doc.embedding is not None else None,
                        "embedded": doc.embedded,
                        "added_at": doc.added_at
                    }
                    corpus_data["documents"].append(doc_dict)

            # Write to file
            filepath = Path(filepath)
            filepath.parent.mkdir(parents=True, exist_ok=True)

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(corpus_data, f, indent=2, ensure_ascii=False)

            logger.info(f"Corpus saved to: {filepath}")
            return True, f"Corpus saved to: {filepath}"

        except Exception as e:
            error_msg = f"Error saving corpus: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    def load_corpus(self, filepath: str) -> Tuple[bool, str]:
        """
        Load corpus from JSON file.

        Args:
            filepath: Path to corpus file

        Returns:
            Tuple of (success, message)
        """
        try:
            filepath = Path(filepath)
            if not filepath.exists():
                return False, f"File not found: {filepath}"

            with open(filepath, 'r', encoding='utf-8') as f:
                corpus_data = json.load(f)

            # Validate format
            if "documents" not in corpus_data:
                return False, "Invalid corpus file format"

            # Load documents
            with self._lock:
                self.documents.clear()

                for doc_dict in corpus_data["documents"]:
                    # Convert embedding list back to numpy array
                    embedding = None
                    if doc_dict.get("embedding") is not None:
                        embedding = np.array(doc_dict["embedding"], dtype=np.float32)

                    doc = CorpusDocument(
                        doc_id=doc_dict["doc_id"],
                        filename=doc_dict["filename"],
                        title=doc_dict["title"],
                        category=doc_dict["category"],
                        source=doc_dict["source"],
                        text_content=doc_dict["text_content"],
                        key_provisions=doc_dict.get("key_provisions", []),
                        escalation_triggers=doc_dict.get("escalation_triggers", []),
                        embedding=embedding,
                        embedded=doc_dict.get("embedded", False),
                        added_at=doc_dict["added_at"]
                    )

                    self.documents[doc.doc_id] = doc

            doc_count = len(self.documents)
            logger.info(f"Loaded {doc_count} documents from: {filepath}")
            return True, f"Loaded {doc_count} documents from corpus"

        except json.JSONDecodeError as e:
            error_msg = f"Invalid JSON in corpus file: {str(e)}"
            logger.error(error_msg)
            return False, error_msg
        except Exception as e:
            error_msg = f"Error loading corpus: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    # ========== Utility Methods ==========

    def get_stats(self) -> Dict[str, Any]:
        """
        Get comprehensive corpus statistics.

        Returns:
            Dictionary with corpus statistics
        """
        with self._lock:
            total = len(self.documents)
            embedded = sum(1 for doc in self.documents.values() if doc.embedded)

            categories = {}
            sources = {}
            total_text_length = 0
            total_provisions = 0
            total_triggers = 0

            for doc in self.documents.values():
                # Count by category
                categories[doc.category] = categories.get(doc.category, 0) + 1

                # Count by source
                sources[doc.source] = sources.get(doc.source, 0) + 1

                # Aggregate stats
                total_text_length += len(doc.text_content)
                total_provisions += len(doc.key_provisions)
                total_triggers += len(doc.escalation_triggers)

            return {
                "total_documents": total,
                "embedded_documents": embedded,
                "not_embedded": total - embedded,
                "embedding_percentage": (embedded / total * 100) if total > 0 else 0,
                "categories": categories,
                "sources": sources,
                "total_text_length": total_text_length,
                "avg_text_length": total_text_length / total if total > 0 else 0,
                "total_key_provisions": total_provisions,
                "total_escalation_triggers": total_triggers
            }


# ========== Public API Helper ==========

_corpus_engine_instance = None


def get_corpus_engine() -> CorpusEngine:
    """
    Get or create the global CorpusEngine instance.

    Returns:
        CorpusEngine instance
    """
    global _corpus_engine_instance
    if _corpus_engine_instance is None:
        _corpus_engine_instance = CorpusEngine()
    return _corpus_engine_instance
